#!/usr/bin/env python3
"""
Service CV - Version complète avec toutes les fonctionnalités
"""

import tempfile
import os
import hashlib
import uuid
from datetime import datetime
from typing import List, Optional, Dict, Any , Tuple
import traceback
from io import BytesIO
from docx import Document
from fastapi.responses import StreamingResponse
from app.services.file_storage import FileStorageService

try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

from app.parsers.text_extractors import TextExtractor
from app.parsers.info_extractors import InfoExtractor
from app.repositories.cv_repository import CVRepository
from app.models.cv_model import CVData, PersonalInfo, CVMetadata, Experience, Formation, LanguageSkill
from bson import ObjectId
from app.repositories.cv_repository import get_cv_collection

class DocumentConverter:
    """Service de conversion de documents"""
    
    def __init__(self):
        self.conversion_methods = []
        
        if DOCX2PDF_AVAILABLE:
            self.conversion_methods.append("docx2pdf")
        if MAMMOTH_AVAILABLE and WEASYPRINT_AVAILABLE:
            self.conversion_methods.append("mammoth_weasyprint")
            
        print(f"📄 Méthodes de conversion disponibles: {self.conversion_methods}")
    
    async def convert_docx_to_pdf(self, docx_content: bytes, filename: str) -> Optional[bytes]:
        """Convertit un fichier DOCX en PDF"""
        try:
            print(f"🔄 Conversion DOCX vers PDF: {filename}")
            
            # Méthode 1: docx2pdf (recommandée sur Windows)
            if "docx2pdf" in self.conversion_methods:
                try:
                    return await self._convert_with_docx2pdf(docx_content, filename)
                except Exception as e:
                    print(f"⚠️ Échec docx2pdf: {e}")
            
            # Méthode 2: mammoth + weasyprint (cross-platform)
            if "mammoth_weasyprint" in self.conversion_methods:
                try:
                    return await self._convert_with_mammoth(docx_content, filename)
                except Exception as e:
                    print(f"⚠️ Échec mammoth: {e}")
            
            print("❌ Aucune méthode de conversion disponible")
            return None
            
        except Exception as e:
            print(f"❌ Erreur conversion {filename}: {e}")
            return None
    
    async def _convert_with_docx2pdf(self, docx_content: bytes, filename: str) -> bytes:
        """Conversion avec docx2pdf"""
        docx_temp = None
        pdf_temp = None
        
        try:
            # Créer fichier DOCX temporaire
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as f:
                f.write(docx_content)
                docx_temp = f.name
            
            # Créer fichier PDF temporaire
            pdf_temp = docx_temp.replace('.docx', '.pdf')
            
            # Conversion
            convert(docx_temp, pdf_temp)
            
            # Lire le PDF généré
            with open(pdf_temp, 'rb') as f:
                pdf_content = f.read()
            
            print(f"✅ Conversion docx2pdf réussie: {len(pdf_content)} bytes")
            return pdf_content
            
        finally:
            # Nettoyer les fichiers temporaires
            for temp_file in [docx_temp, pdf_temp]:
                if temp_file and os.path.exists(temp_file):
                    try:
                        os.unlink(temp_file)
                    except:
                        pass
    
    async def _convert_with_mammoth(self, docx_content: bytes, filename: str) -> bytes:
        """Conversion avec mammoth + weasyprint"""
        try:
            # Convertir DOCX en HTML
            with BytesIO(docx_content) as docx_stream:
                result = mammoth.convert_to_html(docx_stream)
                html_content = result.value
            
            # Créer le HTML complet avec styles
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>{filename}</title>
                <style>
                    body {{ 
                        font-family: Arial, sans-serif; 
                        margin: 20px;
                        line-height: 1.6;
                    }}
                    h1, h2, h3 {{ color: #333; }}
                    p {{ margin-bottom: 10px; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # Convertir HTML en PDF
            html_doc = HTML(string=full_html)
            pdf_bytes = html_doc.write_pdf()
            
            print(f"✅ Conversion mammoth réussie: {len(pdf_bytes)} bytes")
            return pdf_bytes
            
        except Exception as e:
            print(f"❌ Erreur mammoth: {e}")
            raise
    
    def is_conversion_available(self) -> bool:
        """Vérifie si la conversion est disponible"""
        return len(self.conversion_methods) > 0
    
    def get_supported_formats(self) -> list:
        """Retourne les formats supportés pour conversion"""
        if self.is_conversion_available():
            return ['.docx', '.doc']
        return []



class CVService:
    """Service pour la gestion complète des CV"""
    
    def __init__(self):
        self.text_extractor = TextExtractor()
        self.info_extractor = InfoExtractor()
        self.cv_repository = CVRepository()
        self.document_converter = DocumentConverter()
        self.file_storage = FileStorageService()
    
    async def process_cv_file(self, file_content: bytes, filename: str, file_ext: str) -> CVData:
        """Traite un fichier CV complet"""
        print(f"🔄 Début du traitement du fichier: {filename}")
        print(f"📊 Taille du fichier: {len(file_content)} bytes")
        print(f"📄 Extension: {file_ext}")
        
        tmp_file_path = None
        
        try:
            # 1. Créer un fichier temporaire
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
                print(f"📁 Fichier temporaire créé: {tmp_file_path}")
            
            # 2. Extraction du texte
            print("🔍 Extraction du texte...")
            try:
                text = self.text_extractor.extract_text(tmp_file_path)
                if not text or len(text.strip()) == 0:
                    raise Exception("Aucun texte extrait du fichier")
                print(f"📝 Texte extrait: {len(text)} caractères")
            except Exception as e:
                print(f"❌ Erreur extraction texte: {e}")
                raise Exception(f"Impossible d'extraire le texte du fichier: {str(e)}")
            
            # 3. Extraction des informations structurées
            print("🧠 Extraction des informations structurées...")
            try:
                extracted_data = self.info_extractor.extract_all_data(text)
                print("✅ Informations extraites avec succès")
            except Exception as e:
                print(f"❌ Erreur extraction informations: {e}")
                print(f"Traceback: {traceback.format_exc()}")
                raise Exception(f"Erreur lors de l'extraction des informations: {str(e)}")
            
            # 4. Convertir en modèles Pydantic
            print("🔄 Conversion en modèles Pydantic...")
            
            # PersonalInfo avec validation
            try:
                personal_info_data = extracted_data.get('informations_personnelles', {})
                personal_info = PersonalInfo(
                    nom=personal_info_data.get('nom') if personal_info_data.get('nom') != "Non trouvé" else None,
                    email=personal_info_data.get('email') if personal_info_data.get('email') != "Non trouvé" else None,
                    telephone=personal_info_data.get('telephone') if personal_info_data.get('telephone') != "Non trouvé" else None,
                    adresse=personal_info_data.get('adresse') if personal_info_data.get('adresse') != "Non trouvé" else None
                )
                print(f"✅ PersonalInfo créé: {personal_info.nom}")
            except Exception as e:
                print(f"⚠️ Erreur création PersonalInfo: {e}")
                personal_info = PersonalInfo()
            
            # CVMetadata avec validation
            try:
                metadata_raw = extracted_data.get('metadonnees', {})
                metadata = CVMetadata(
                    nombre_mots=metadata_raw.get('nombre_mots', len(text.split()) if text else 0),
                    date_extraction=metadata_raw.get('date_extraction', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
                    apercu_texte=metadata_raw.get('apercu_texte', text[:200] if text else ""),
                    taille_fichier_kb=round(len(file_content) / 1024, 2)
                )
                print(f"✅ CVMetadata créé: {metadata.nombre_mots} mots")
            except Exception as e:
                print(f"⚠️ Erreur création CVMetadata: {e}")
                metadata = CVMetadata(
                    nombre_mots=len(text.split()) if text else 0,
                    date_extraction=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    apercu_texte=text[:200] if text else "",
                    taille_fichier_kb=round(len(file_content) / 1024, 2)
                )
            
            # Convertir les expériences
            experiences = []
            experiences_data = extracted_data.get('experience_professionnelle', [])
            print(f"🔄 Traitement de {len(experiences_data)} expériences...")
            
            for i, exp_data in enumerate(experiences_data):
                try:
                    if isinstance(exp_data, dict):
                        experience = Experience(
                            periode=exp_data.get('periode', ''),
                            poste=exp_data.get('poste', ''),
                            entreprise=exp_data.get('entreprise', ''),
                            description=exp_data.get('description', '')
                        )
                        experiences.append(experience)
                        print(f"✅ Expérience {i+1} créée: {experience.poste}")
                except Exception as e:
                    print(f"⚠️ Erreur création Experience {i}: {e}")
                    continue
            
            # Convertir les formations
            formations = []
            formations_data = extracted_data.get('formations_academiques', [])
            print(f"🔄 Traitement de {len(formations_data)} formations...")
            
            for i, form_data in enumerate(formations_data):
                try:
                    if isinstance(form_data, dict):
                        formation = Formation(
                            annee=form_data.get('annee', ''),
                            diplome=form_data.get('diplome', ''),
                            etablissement=form_data.get('etablissement', ''),
                            mention=form_data.get('mention', '')
                        )
                        formations.append(formation)
                        print(f"✅ Formation {i+1} créée: {formation.diplome}")
                except Exception as e:
                    print(f"⚠️ Erreur création Formation {i}: {e}")
                    continue
            
            # Convertir les langues
            langues = []
            langues_data = extracted_data.get('competences_linguistiques', [])
            print(f"🔄 Traitement de {len(langues_data)} langues...")
            
            for i, lang_data in enumerate(langues_data):
                try:
                    if isinstance(lang_data, dict) and lang_data.get('langue'):
                        langue = LanguageSkill(
                            langue=lang_data.get('langue', ''),
                            niveau=lang_data.get('niveau', 'Non spécifié')
                        )
                        langues.append(langue)
                        print(f"✅ Langue {i+1} créée: {langue.langue} - {langue.niveau}")
                except Exception as e:
                    print(f"⚠️ Erreur création LanguageSkill {i}: {e}")
                    continue
            
            # 5. Créer l'objet CVData
            print("📦 Création de l'objet CVData...")
            try:
                cv_id = str(uuid.uuid4())
                
                competences_techniques = extracted_data.get('competences_techniques', [])
                if not isinstance(competences_techniques, list):
                    competences_techniques = []
                
                certifications = extracted_data.get('certifications', [])
                if not isinstance(certifications, list):
                    certifications = []
                
                cv_data = CVData(
                    id=cv_id,
                    informations_personnelles=personal_info,
                    competences_techniques=competences_techniques,
                    experience_professionnelle=experiences,
                    formations_academiques=formations,
                    certifications=certifications,
                    competences_linguistiques=langues,
                    type_document=extracted_data.get('type_document', 'CV'),
                    metadonnees=metadata,
                    nlp_enrichment=extracted_data.get('nlp_enrichment'),
                    filename_original=filename,
                    file_hash=self._calculate_file_hash(file_content),
                    status="completed",
                    created_at=datetime.now(),
                    updated_at=datetime.now()
                )
                
                print(f"✅ CVData créé avec l'ID: {cv_id}")
                
            except Exception as e:
                print(f"❌ Erreur création CVData: {e}")
                print(f"Traceback: {traceback.format_exc()}")
                raise Exception(f"Erreur lors de la création de l'objet CV: {str(e)}")
            
            # 6. Stocker le fichier original - NOUVEAU
            print("💾 Stockage du fichier original...")
            try:
                file_stored = await self.file_storage.store_file(
                    file_content, cv_id, filename
                )
                if file_stored:
                    print(f"✅ Fichier original stocké: {cv_id}")
                else:
                    print("⚠️ Échec du stockage du fichier original")
            except Exception as e:
                print(f"❌ Erreur stockage fichier: {e}")
            
            # 6. Sauvegarder en base
            print("💾 Tentative de sauvegarde en base de données...")
            try:
                saved_cv = await self.cv_repository.create_cv(cv_data)
                if saved_cv:
                    print(f"✅ CV sauvegardé en base avec succès: {cv_id}")
                    return saved_cv
                else:
                    print("⚠️ La sauvegarde a retourné None")
                    return cv_data
            except Exception as e:
                print(f"❌ Impossible de sauvegarder en base: {e}")
                print("📝 Le CV sera retourné sans sauvegarde")
                cv_data.status = "not_saved"
                return cv_data
            
        except Exception as e:
            print(f"❌ Erreur lors du traitement: {e}")
            print(f"Traceback complet: {traceback.format_exc()}")
            raise Exception(f"Erreur lors du traitement du CV: {str(e)}")
            
        finally:
            # Nettoyer le fichier temporaire
            if tmp_file_path and os.path.exists(tmp_file_path):
                try:
                    os.unlink(tmp_file_path)
                    print("🗑️ Fichier temporaire supprimé")
                except Exception as e:
                    print(f"⚠️ Erreur suppression fichier temporaire: {e}")
    
    async def update_cv(self, cv_data: CVData) -> Optional[CVData]:
        """Met à jour un CV existant - VERSION CORRIGÉE"""
        try:
            print(f"🔄 Mise à jour du CV: {cv_data.id}")
            
            # Vérifier que le CV existe
            existing_cv = await self.cv_repository.get_cv_by_id(cv_data.id)
            if not existing_cv:
                print(f"❌ CV non trouvé pour mise à jour: {cv_data.id}")
                return None
            
            # S'assurer que les dates sont conservées
            if not cv_data.created_at and existing_cv.created_at:
                cv_data.created_at = existing_cv.created_at
            
            cv_data.updated_at = datetime.now()
            
            # Utiliser le repository pour la mise à jour
            updated_cv = await self.cv_repository.update_cv(cv_data)
            
            if updated_cv:
                updated_cv.status = "completed"
                print(f"✅ CV mis à jour: {cv_data.id}")
            else:
                print(f"❌ Échec mise à jour CV: {cv_data.id}")
                
            return updated_cv
            
        except Exception as e:
            print(f"❌ Erreur mise à jour CV {cv_data.id}: {e}")
            print(f"Traceback: {traceback.format_exc()}")
            return None

    async def update_cv_fields(self, cv_id: str, updates: Dict[str, Any]) -> Optional[CVData]:
        """Met à jour des champs spécifiques d'un CV - VERSION CORRIGÉE"""
        try:
            print(f"🔄 Mise à jour partielle CV: {cv_id}")
            print(f"📝 Champs: {list(updates.keys())}")
            
            # Récupérer le CV existant
            existing_cv = await self.cv_repository.get_cv_by_id(cv_id)
            if not existing_cv:
                print(f"❌ CV non trouvé: {cv_id}")
                return None
            
            # Convertir en dictionnaire
            cv_dict = existing_cv.dict()
            
            # Appliquer les mises à jour récursivement pour les sous-structures
            def update_dict_recursive(original, updates):
                for key, value in updates.items():
                    if isinstance(value, dict) and key in original and isinstance(original[key], dict):
                        update_dict_recursive(original[key], value)
                    else:
                        original[key] = value
            
            update_dict_recursive(cv_dict, updates)
            cv_dict["updated_at"] = datetime.now()
            
            # Recréer l'objet CVData
            updated_cv_data = CVData(**cv_dict)
            
            # Sauvegarder
            result = await self.cv_repository.update_cv(updated_cv_data)
            
            if result:
                print(f"✅ Champs mis à jour: {cv_id}")
            else:
                print(f"❌ Échec mise à jour champs: {cv_id}")
                
            return result
            
        except Exception as e:
            print(f"❌ Erreur mise à jour champs CV {cv_id}: {e}")
            print(f"Traceback: {traceback.format_exc()}")
            return None
    
    async def extract_text_only(self, file_content: bytes, file_ext: str) -> str:
        """Extrait seulement le texte d'un fichier"""
        print(f"📄 Extraction de texte pour fichier {file_ext}")
        
        tmp_file_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            text = self.text_extractor.extract_text(tmp_file_path)
            print(f"📝 Texte extrait: {len(text) if text else 0} caractères")
            return text or ""
            
        except Exception as e:
            print(f"❌ Erreur extraction texte: {e}")
            return ""
            
        finally:
            if tmp_file_path and os.path.exists(tmp_file_path):
                try:
                    os.unlink(tmp_file_path)
                except Exception as e:
                    print(f"⚠️ Erreur suppression fichier temporaire: {e}")
    
    async def get_all_cvs(self) -> List[CVData]:
        """Récupère tous les CV"""
        print("📋 Récupération de tous les CV...")
        try:
            cvs = await self.cv_repository.get_all_cvs()
            print(f"✅ {len(cvs)} CV(s) récupéré(s)")
            return cvs
        except Exception as e:
            print(f"❌ Erreur récupération CV: {e}")
            return []
    
    async def get_cv_by_id(self, cv_id: str) -> Optional[CVData]:
        """Récupère un CV par ID (stricte puis flexible)"""
        try:
            print(f"🔍 Service: Recherche CV: {cv_id}")
            if not cv_id:
                print("❌ ID manquant")
                return None

            # Recherche stricte par id
            cv = await self.cv_repository.get_cv_by_id(cv_id)
            if cv:
                print(f"✅ Service: CV trouvé (stricte): {cv_id}")
                return cv

            # Si non trouvé, recherche flexible (id ou _id)
            cv = await self.cv_repository.get_cv_by_any_id(cv_id)
            if cv:
                print(f"✅ Service: CV trouvé (flexible): {cv_id}")
                return cv

            print(f"❌ Service: CV non trouvé: {cv_id}")
            return None

        except Exception as e:
            print(f"❌ Erreur service récupération CV {cv_id}: {e}")
            print(f"Traceback: {traceback.format_exc()}")
            return None
    
    async def delete_cv(self, cv_id: str) -> bool:
        """Supprime un CV et son fichier associé - VERSION MISE À JOUR"""
        print(f"🗑️ Suppression du CV: {cv_id}")
        try:
            # Récupérer les infos du CV avant suppression
            cv_data = await self.cv_repository.get_cv_by_id(cv_id)
            
            # Supprimer de la base de données
            result = await self.cv_repository.delete_cv(cv_id)
            
            # Si succès, supprimer aussi le fichier stocké
            if result and cv_data and cv_data.filename_original:
                try:
                    await self.file_storage.delete_file(cv_id, cv_data.filename_original)
                    print(f"✅ Fichier supprimé: {cv_id}")
                except Exception as e:
                    print(f"⚠️ Erreur suppression fichier: {e}")
                    # Ne pas faire échouer la suppression DB si fichier pas supprimé
            
            if result:
                print(f"✅ CV supprimé: {cv_id}")
            else:
                print(f"❌ Échec suppression CV: {cv_id}")
            return result
        except Exception as e:
            print(f"❌ Erreur suppression CV {cv_id}: {e}")
            return False
    
    async def search_cvs_by_skills(self, skills: List[str]) -> List[CVData]:
        """Recherche des CV par compétences"""
        print(f"🔍 Recherche CV par compétences: {skills}")
        try:
            cvs = await self.cv_repository.search_by_skills(skills)
            print(f"✅ {len(cvs)} CV(s) trouvé(s) avec ces compétences")
            return cvs
        except Exception as e:
            print(f"❌ Erreur recherche CV par compétences: {e}")
            return []
    
    async def update_cv_status(self, cv_id: str, status: str) -> bool:
        """Met à jour le statut d'un CV"""
        print(f"🔄 Mise à jour statut CV {cv_id}: {status}")
        try:
            result = await self.cv_repository.update_cv_status(cv_id, status)
            if result:
                print(f"✅ Statut mis à jour: {cv_id}")
            return result
        except Exception as e:
            print(f"❌ Erreur mise à jour statut CV {cv_id}: {e}")
            return False
    
    def _calculate_file_hash(self, file_content: bytes) -> str:
        """Calcule le hash MD5 du fichier"""
        try:
            hash_md5 = hashlib.md5(file_content).hexdigest()
            print(f"🔒 Hash calculé: {hash_md5[:8]}...")
            return hash_md5
        except Exception as e:
            print(f"❌ Erreur calcul hash: {e}")
            return str(uuid.uuid4())  # Fallback
    
    async def validate_cv_data(self, cv_data: CVData) -> bool:
        """Valide la cohérence des données d'un CV"""
        print(f"✅ Validation des données CV: {cv_data.id}")
        try:
            if not cv_data.id:
                print("❌ ID manquant")
                return False
            
            if not cv_data.filename_original:
                print("❌ Nom de fichier manquant")
                return False
            
            if not cv_data.file_hash:
                print("❌ Hash de fichier manquant")
                return False
            
            print("✅ Données CV valides")
            return True
            
        except Exception as e:
            print(f"❌ Erreur validation CV: {e}")
            return False
    
    async def get_cv_statistics(self) -> Dict[str, Any]:
        """Récupère les statistiques des CV"""
        try:
            print("📊 Calcul des statistiques CV...")
            cvs = await self.get_all_cvs()
            
            total_cvs = len(cvs)
            completed_cvs = len([cv for cv in cvs if cv.status == "completed"])
            processing_cvs = len([cv for cv in cvs if cv.status == "processing"])
            error_cvs = len([cv for cv in cvs if cv.status == "error"])
            
            # Compétences les plus communes
            all_skills = []
            for cv in cvs:
                if cv.competences_techniques:
                    all_skills.extend(cv.competences_techniques)
            
            from collections import Counter
            top_skills = Counter(all_skills).most_common(10)
            
            stats = {
                "total_cvs": total_cvs,
                "completed_cvs": completed_cvs,
                "processing_cvs": processing_cvs,
                "error_cvs": error_cvs,
                "success_rate": round((completed_cvs / total_cvs * 100) if total_cvs > 0 else 0, 2),
                "top_skills": [{"skill": skill, "count": count} for skill, count in top_skills],
                "last_updated": datetime.now().isoformat()
            }
            
            print(f"✅ Statistiques calculées: {total_cvs} CV(s)")
            return stats
            
        except Exception as e:
            print(f"❌ Erreur calcul statistiques: {e}")
            return {
                "total_cvs": 0,
                "completed_cvs": 0,
                "processing_cvs": 0,
                "error_cvs": 0,
                "success_rate": 0,
                "top_skills": [],
                "last_updated": datetime.now().isoformat()
            }
    
    async def export_cv_onetech(self, cv_id: str):
        """Exporte un CV au format OneTech (.docx)"""
        try:
            print(f"📤 Service: Export OneTech (Word) pour CV: {cv_id}")
            
            # Récupération du CV
            cv_data = await self.get_cv_by_id(cv_id)
            if not cv_data:
                print(f"❌ CV non trouvé pour export OneTech: {cv_id}")
                return None

            # Création du document Word
            doc = Document()
            doc.add_heading("Curriculum Vitae - OneTech", level=0)

            # Infos personnelles
            doc.add_heading("Informations personnelles", level=1)
            doc.add_paragraph(f"Nom : {cv_data.informations_personnelles.nom if cv_data.informations_personnelles else 'Non spécifié'}")
            doc.add_paragraph(f"Email : {cv_data.informations_personnelles.email if cv_data.informations_personnelles else 'Non spécifié'}")
            doc.add_paragraph(f"Téléphone : {cv_data.informations_personnelles.telephone if cv_data.informations_personnelles else 'Non spécifié'}")
            doc.add_paragraph(f"Adresse : {cv_data.informations_personnelles.adresse if cv_data.informations_personnelles else 'Non spécifié'}")

            # Compétences techniques
            doc.add_heading("Compétences techniques", level=1)
            for skill in (cv_data.competences_techniques or []):
                doc.add_paragraph(f"- {skill}", style="List Bullet")

            # Expérience professionnelle
            doc.add_heading("Expérience professionnelle", level=1)
            for exp in (cv_data.experience_professionnelle or []):
                doc.add_paragraph(f"{exp.periode} - {exp.poste} @ {exp.entreprise}")
                doc.add_paragraph(exp.description, style="Intense Quote")

            # Formations académiques
            doc.add_heading("Formations académiques", level=1)
            for form in (cv_data.formations_academiques or []):
                doc.add_paragraph(f"{form.annee} - {form.diplome} ({form.etablissement}) - Mention: {form.mention}")

            # Langues
            doc.add_heading("Langues", level=1)
            for lang in (cv_data.competences_linguistiques or []):
                doc.add_paragraph(f"{lang.langue} : {lang.niveau}")

            # Métadonnées
            doc.add_heading("Métadonnées", level=1)
            doc.add_paragraph(f"Nombre de mots: {cv_data.metadonnees.nombre_mots if cv_data.metadonnees else 0}")
            doc.add_paragraph(f"Date extraction: {cv_data.metadonnees.date_extraction if cv_data.metadonnees else 'N/A'}")
            doc.add_paragraph(f"Taille fichier (KB): {cv_data.metadonnees.taille_fichier_kb if cv_data.metadonnees else 0}")

            # Sauvegarde en mémoire
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            print(f"✅ Export OneTech Word réussi pour CV: {cv_id}")
            return StreamingResponse(
                buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=CV_{cv_id}_OneTech.docx"}
            )

        except Exception as e:
            print(f"❌ Erreur export OneTech CV {cv_id}: {e}")
            print(f"Traceback: {traceback.format_exc()}")
            raise Exception(f"Erreur lors de l'export OneTech: {str(e)}")
    async def export_cv_json(self, cv_id: str) -> Optional[dict]:
        """Retourne le CV brut en JSON tel qu'il est stocké en base"""
        collection = get_cv_collection()  # <-- au lieu de self._get_collection()

        query = {"$or": [{"id": cv_id}]}
        if ObjectId.is_valid(cv_id):
            query["$or"].append({"_id": ObjectId(cv_id)})

        cv_doc = await collection.find_one(query)
        if not cv_doc:
            return None

        # Normalisation de l'id
        if isinstance(cv_doc.get("_id"), ObjectId):
            cv_doc["id"] = str(cv_doc["_id"])
        elif "_id" in cv_doc:
            cv_doc["id"] = cv_doc["_id"]

        cv_doc.pop("_id", None)
        return cv_doc
    
    async def check_and_process_cv(self, file_content: bytes, filename: str, file_ext: str) -> tuple[Optional[CVData], bool]:
        """Vérifie si le CV existe déjà et le traite si nécessaire - VERSION CORRIGÉE"""
        try:
            # Calculer le hash du fichier
            file_hash = self._calculate_file_hash(file_content)
            
            # Vérifier si un CV avec ce hash existe déjà
            existing_cv = await self.cv_repository.check_duplicate_hash(file_hash)
            
            if existing_cv:
                print(f"⚠️ CV déjà existant: {file_hash[:8]}...")
                
                # NOUVEAU: Vérifier si le fichier associé existe
                file_exists = self.file_storage.file_exists(existing_cv.id, existing_cv.filename_original or filename)
                
                if not file_exists:
                    print(f"⚠️ Fichier manquant pour CV existant, stockage du nouveau fichier...")
                    try:
                        # Stocker le fichier pour le CV existant
                        file_stored = await self.file_storage.store_file(
                            file_content, existing_cv.id, existing_cv.filename_original or filename, replace_existing=True
                        )
                        if file_stored:
                            print(f"✅ Fichier stocké pour CV existant: {existing_cv.id}")
                            
                            # Mettre à jour le nom du fichier si nécessaire
                            if not existing_cv.filename_original or existing_cv.filename_original != filename:
                                await self.update_cv_fields(existing_cv.id, {
                                    "filename_original": filename,
                                    "file_hash": file_hash,
                                    "updated_at": datetime.now()
                                })
                        else:
                            print(f"❌ Échec stockage fichier pour CV existant: {existing_cv.id}")
                    except Exception as e:
                        print(f"❌ Erreur stockage fichier CV existant: {e}")
                
                return existing_cv, True  # CV existant, doublon = True
            
            # Si pas de doublon, traiter normalement
            cv_data = await self.process_cv_file(file_content, filename, file_ext)
            return cv_data, False
            
        except Exception as e:
            print(f"❌ Erreur vérification doublon: {e}")
            raise

    # NOUVELLE MÉTHODE pour remplacer un CV avec gestion du fichier
    async def replace_cv_with_file(self, cv_id: str, file_content: bytes, filename: str, file_ext: str) -> Optional[CVData]:
        """Remplace un CV existant avec un nouveau fichier - VERSION CORRIGÉE"""
        try:
            print(f"🔄 Remplacement du CV avec fichier: {cv_id}")
            
            # Vérifier que le CV existe
            existing_cv = await self.get_cv_by_id(cv_id)
            if not existing_cv:
                print(f"❌ CV à remplacer non trouvé: {cv_id}")
                return None
            
            # Traiter le nouveau fichier
            new_cv_data = await self.process_cv_file(file_content, filename, file_ext)
            
            # Conserver l'ID original et certaines métadonnées
            new_cv_data.id = cv_id
            new_cv_data.created_at = existing_cv.created_at  # Conserver la date de création
            new_cv_data.updated_at = datetime.now()
            
            # Stocker le nouveau fichier (remplacer l'ancien)
            print("💾 Remplacement du fichier original...")
            try:
                file_stored = await self.file_storage.store_file(
                    file_content, cv_id, filename, replace_existing=True
                )
                if file_stored:
                    print(f"✅ Nouveau fichier stocké: {cv_id}")
                else:
                    print("⚠️ Échec du remplacement du fichier")
            except Exception as e:
                print(f"❌ Erreur remplacement fichier: {e}")
            
            # Sauvegarder les nouvelles données
            updated_cv = await self.update_cv(new_cv_data)
            
            if updated_cv:
                print(f"✅ CV et fichier remplacés: {cv_id}")
            else:
                print(f"❌ Échec remplacement CV: {cv_id}")
                
            return updated_cv
            
        except Exception as e:
            print(f"❌ Erreur remplacement CV avec fichier {cv_id}: {e}")
            print(f"Traceback: {traceback.format_exc()}")
            return None


    async def export_cv_text(self, cv_id: str) -> Optional[str]:
        """Exporte un CV au format texte"""
        try:
            print(f"📤 Export texte pour CV: {cv_id}")
            
            cv_data = await self.get_cv_by_id(cv_id)
            if not cv_data:
                return None
            
            # Construire le texte formaté
            text_parts = []
            
            # Informations personnelles
            text_parts.append("=== INFORMATIONS PERSONNELLES ===")
            if cv_data.informations_personnelles.nom:
                text_parts.append(f"Nom: {cv_data.informations_personnelles.nom}")
            if cv_data.informations_personnelles.email:
                text_parts.append(f"Email: {cv_data.informations_personnelles.email}")
            if cv_data.informations_personnelles.telephone:
                text_parts.append(f"Téléphone: {cv_data.informations_personnelles.telephone}")
            if cv_data.informations_personnelles.adresse:
                text_parts.append(f"Adresse: {cv_data.informations_personnelles.adresse}")
            
            # Compétences techniques
            if cv_data.competences_techniques:
                text_parts.append("\n=== COMPÉTENCES TECHNIQUES ===")
                text_parts.extend(cv_data.competences_techniques)
            
            # Expérience professionnelle
            if cv_data.experience_professionnelle:
                text_parts.append("\n=== EXPÉRIENCE PROFESSIONNELLE ===")
                for exp in cv_data.experience_professionnelle:
                    text_parts.append(f"\n{exp.periode} - {exp.poste}")
                    text_parts.append(f"Entreprise: {exp.entreprise}")
                    if exp.description:
                        text_parts.append(f"Description: {exp.description}")
        
            # Formations
            if cv_data.formations_academiques:
                text_parts.append("\n=== FORMATIONS ACADÉMIQUES ===")
                for form in cv_data.formations_academiques:
                    text_parts.append(f"\n{form.annee} - {form.diplome}")
                    text_parts.append(f"Établissement: {form.etablissement}")
                    if form.mention:
                        text_parts.append(f"Mention: {form.mention}")
            
            # Langues
            if cv_data.competences_linguistiques:
                text_parts.append("\n=== COMPÉTENCES LINGUISTIQUES ===")
                for lang in cv_data.competences_linguistiques:
                    text_parts.append(f"{lang.langue}: {lang.niveau}")
            
            # Certifications
            if cv_data.certifications:
                text_parts.append("\n=== CERTIFICATIONS ===")
                text_parts.extend(cv_data.certifications)
            
            result_text = "\n".join(text_parts)
            print(f"✅ Export texte réussi pour CV: {cv_id}")
            return result_text
            
        except Exception as e:
            print(f"❌ Erreur export texte CV {cv_id}: {e}")
            return None

            
        except Exception as e:
            print(f"❌ Erreur export texte CV {cv_id}: {e}")
            return None

    async def get_document_for_preview(self, cv_id: str) -> Tuple[Optional[bytes], str, str]:
        """
        Récupère le document pour aperçu, avec conversion si nécessaire - VERSION CORRIGÉE
        Returns: (content, content_type, original_filename)
        """
        try:
            print(f"📄 Récupération document pour aperçu: {cv_id}")
            
            # Récupérer le CV
            cv_data = await self.get_cv_by_id(cv_id)
            if not cv_data or not cv_data.filename_original:
                print(f"❌ CV ou nom de fichier non trouvé: {cv_id}")
                return None, "", ""
            
            # Récupérer le contenu du fichier original
            original_content = await self._get_original_file_content(cv_id)
            if not original_content:
                print(f"❌ Contenu fichier non disponible: {cv_id}")
                return None, "", ""
            
            original_filename = cv_data.filename_original
            file_ext = os.path.splitext(original_filename)[1].lower()
            
            print(f"📄 Type de fichier détecté: {file_ext}")
            
            # Si c'est un PDF, retourner directement
            if file_ext == '.pdf':
                print(f"📄 Retour PDF direct: {original_filename}")
                return original_content, "application/pdf", original_filename
            
            # Si c'est un DOCX et conversion disponible
            if file_ext in ['.docx', '.doc'] and self.document_converter.is_conversion_available():
                print(f"🔄 Tentative de conversion: {original_filename}")
                pdf_content = await self.document_converter.convert_docx_to_pdf(
                    original_content, original_filename
                )
                
                if pdf_content:
                    print(f"✅ Conversion réussie: {original_filename} -> PDF")
                    return pdf_content, "application/pdf", f"{original_filename}_converted.pdf"
                else:
                    print(f"❌ Conversion échouée, retour fichier original: {original_filename}")
                    # Retourner le fichier original même si conversion échoue
                    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    return original_content, content_type, original_filename
            
            # Pour autres formats, retourner tel quel
            print(f"📄 Retour fichier original sans conversion: {original_filename}")
            return original_content, "application/octet-stream", original_filename
            
        except Exception as e:
            print(f"❌ Erreur récupération document {cv_id}: {e}")
            print(f"Traceback: {traceback.format_exc()}")
            return None, "", ""
    
    async def _get_original_file_content(self, cv_id: str) -> Optional[bytes]:
        """Récupère le contenu du fichier original - VERSION CORRIGÉE"""
        try:
            # Récupérer les métadonnées du CV pour avoir le nom de fichier original
            cv_data = await self.get_cv_by_id(cv_id)
            if not cv_data or not cv_data.filename_original:
                print(f"❌ CV ou nom de fichier non trouvé: {cv_id}")
                return None
            
            # Utiliser le service de stockage pour récupérer le fichier
            content = await self.file_storage.get_file_content(cv_id, cv_data.filename_original)
            if content:
                print(f"✅ Contenu fichier récupéré: {cv_id} ({len(content)} bytes)")
                return content
            else:
                print(f"❌ Contenu fichier non trouvé: {cv_id}")
                return None
                
        except Exception as e:
            print(f"❌ Erreur récupération fichier {cv_id}: {e}")
            return None

